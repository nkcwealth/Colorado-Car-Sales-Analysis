from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_word_report():
    # Create a new Document
    doc = Document()
    
    # Title
    title = doc.add_heading('Colorado Motor Vehicle Sales Analysis Project Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph('Data Analysis and Visualization Project')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Date: April 2025').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Executive Summary
    doc.add_heading('Executive Summary', 1)
    doc.add_paragraph('This report presents a comprehensive analysis of car sales data in Colorado, focusing on key metrics including sales trends, pricing patterns, vehicle preferences, and geographic distribution. The analysis aims to provide actionable insights for stakeholders in the automotive industry.')
    
    # Add Market Overview
    doc.add_heading('Market Overview', 1)
    doc.add_picture('reports/figures/engine_type_distribution.png', width=Inches(6))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Figure 1: Engine Type Distribution')
    
    overview_stats = doc.add_paragraph()
    overview_stats.add_run('\nKey Statistics:\n').bold = True
    overview_stats.add_run('''
• Total Vehicles Analyzed: 20
• Average Price: $42,665
• Price Range: $25,000 - $62,000''')
    
    # Sales Distribution
    doc.add_heading('Sales Distribution', 1)
    doc.add_picture('reports/figures/sales_type_distribution.png', width=Inches(6))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Figure 2: Sales Type Distribution')
    
    # Geographic Analysis
    doc.add_heading('Geographic Analysis', 1)
    doc.add_picture('reports/figures/sales_by_location.png', width=Inches(6))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Figure 3: Sales by Location')
    
    # Price Analysis
    doc.add_heading('Price Analysis', 1)
    doc.add_picture('reports/figures/price_distribution.png', width=Inches(6))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Figure 4: Price Distribution')
    
    # Vehicle Analysis
    doc.add_heading('Vehicle Analysis', 1)
    doc.add_picture('reports/figures/avg_price_by_make.png', width=Inches(6))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Figure 5: Average Price by Make')
    
    # Price vs Mileage Analysis
    doc.add_heading('Price vs Mileage Analysis', 1)
    doc.add_picture('reports/figures/price_vs_mileage.png', width=Inches(6))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Figure 6: Price vs Mileage Relationship')
    
    # Recommendations
    doc.add_heading('Recommendations', 1)
    
    # Inventory Management
    doc.add_heading('Inventory Management', 2)
    inventory = doc.add_paragraph()
    inventory.add_run('Recommended Inventory Mix:\n').bold = True
    inventory.add_run('''
• 45% Gas vehicles
• 30% Electric vehicles
• 25% Hybrid vehicles

''')
    inventory.add_run('Price Range Focus:\n').bold = True
    inventory.add_run('''
• Entry-level: $25,000 - $35,000
• Mid-range: $35,000 - $45,000
• Premium: $45,000+''')
    
    # Geographic Strategy
    doc.add_heading('Geographic Strategy', 2)
    geo = doc.add_paragraph()
    geo.add_run('Priority Locations:\n').bold = True
    geo.add_run('''
• Denver
• Colorado Springs
• Fort Collins''')
    
    # Technical Details
    doc.add_heading('Technical Implementation', 1)
    tech = doc.add_paragraph()
    tech.add_run('Tools and Technologies:\n').bold = True
    tech.add_run('''
• Python 3.13.3
• pandas 2.2.0+
• matplotlib 3.8.0+
• seaborn 0.13.0+
• numpy 1.26.0+''')
    
    # Conclusions
    doc.add_heading('Conclusions', 1)
    conclusions = doc.add_paragraph()
    conclusions.add_run('Key Takeaways:\n').bold = True
    conclusions.add_run('''
1. Strong market for electric vehicles
2. Geographic variation in sales patterns
3. Price-mileage relationship significance
4. Registration status impact''')
    
    # Save the document
    doc.save('reports/Colorado_Car_Sales_Analysis_Report.docx')
    print("Word document has been created successfully at 'reports/Colorado_Car_Sales_Analysis_Report.docx'")

if __name__ == "__main__":
    create_word_report() 